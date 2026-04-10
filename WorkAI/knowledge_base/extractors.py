"""Knowledge source extractors for markdown, DOCX, and PDF."""

from __future__ import annotations

import json
from pathlib import Path

from WorkAI.knowledge_base.models import KnowledgeArticleDocument

SUPPORTED_KNOWLEDGE_EXTENSIONS: tuple[str, ...] = (".md", ".docx", ".pdf")


def parse_article_document(path: Path) -> KnowledgeArticleDocument:
    """Parse one knowledge source document by extension."""

    suffix = path.suffix.lower()
    if suffix == ".md":
        return parse_markdown_article(path)
    if suffix == ".docx":
        return parse_docx_article(path)
    if suffix == ".pdf":
        return parse_pdf_article(path)

    raise ValueError(f"Unsupported knowledge source extension: {suffix}")


def parse_markdown_article(path: Path) -> KnowledgeArticleDocument:
    """Parse one markdown file into article payload."""

    content = path.read_text(encoding="utf-8")
    metadata, markdown_body = _split_frontmatter(content)

    title, body = _extract_title_and_body(markdown_body, fallback_title=path.stem)
    tags = _extract_tags(metadata)

    return KnowledgeArticleDocument(
        source_path=str(path),
        title=title,
        body=body,
        tags=tags,
    )


def parse_docx_article(path: Path) -> KnowledgeArticleDocument:
    """Parse one DOCX file into article payload."""

    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise RuntimeError("python-docx is required for .docx knowledge sources") from exc

    doc = Document(str(path))
    lines = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    body_text = "\n".join(lines).strip()

    title = path.stem
    if lines:
        first = lines[0]
        if first.startswith("# "):
            title = first[2:].strip() or path.stem
            body_text = "\n".join(lines[1:]).strip()
        else:
            title = first
            body_text = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""

    if body_text == "":
        body_text = "\n".join(lines).strip()

    return KnowledgeArticleDocument(
        source_path=str(path),
        title=title,
        body=body_text,
        tags=[],
    )


def parse_pdf_article(path: Path) -> KnowledgeArticleDocument:
    """Parse one PDF file into article payload."""

    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise RuntimeError("pypdf is required for .pdf knowledge sources") from exc

    reader = PdfReader(str(path))
    chunks: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        stripped = text.strip()
        if stripped:
            chunks.append(stripped)

    body = "\n\n".join(chunks).strip()
    meta_title = None
    if reader.metadata is not None:
        meta_title = reader.metadata.get("/Title")

    title = str(meta_title).strip() if meta_title else ""
    if title == "":
        title = path.stem

    return KnowledgeArticleDocument(
        source_path=str(path),
        title=title,
        body=body,
        tags=[],
    )


def _split_frontmatter(content: str) -> tuple[dict[str, object], str]:
    text = content.lstrip("\ufeff")
    if not text.startswith("---\n"):
        return {}, content

    end_marker = "\n---\n"
    end_index = text.find(end_marker, 4)
    if end_index == -1:
        return {}, content

    raw_meta = text[4:end_index]
    body = text[end_index + len(end_marker) :]
    metadata: dict[str, object] = {}

    for line in raw_meta.splitlines():
        stripped = line.strip()
        if not stripped or ":" not in stripped:
            continue

        key, value = stripped.split(":", 1)
        metadata[key.strip().lower()] = value.strip()

    return metadata, body


def _extract_title_and_body(markdown_body: str, *, fallback_title: str) -> tuple[str, str]:
    lines = markdown_body.splitlines()
    title = fallback_title
    body_lines = lines

    for index, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("# "):
            candidate = stripped[2:].strip()
            if candidate:
                title = candidate
            body_lines = lines[:index] + lines[index + 1 :]
            break

    body = "\n".join(body_lines).strip()
    return title, body


def _extract_tags(metadata: dict[str, object]) -> list[str]:
    raw = metadata.get("tags")
    if raw is None:
        return []

    if isinstance(raw, str):
        candidate = raw.strip()
        if candidate.startswith("[") and candidate.endswith("]"):
            try:
                parsed = json.loads(candidate)
                if isinstance(parsed, list):
                    return [str(item).strip() for item in parsed if str(item).strip()]
            except json.JSONDecodeError:
                pass

        return [part.strip() for part in candidate.split(",") if part.strip()]

    return []
